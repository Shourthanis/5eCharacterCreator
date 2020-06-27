#5th Edition Character Creator
#Shourthanis
#Version 1.0.0

"""
To Do List:

Spellcasting for class display
Spellcasting for export

random overGen?

cleanUp

"""

#import block
from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from tkinter import font
import json
import random #random.randint(low, high)
import math
import docx

#main window initialization and perameters
root = Tk()
root.title("5eCC")

#File load block for races, classes, and backgrounds
readFile = open('Backgrounds.json', "r")
backgroundData = json.load(readFile)
readFile.close()
readFile = open('Races.json', "r")
raceData = json.load(readFile)
readFile.close()
readFile = open('Classes.json', "r")
classData = json.load(readFile)
readFile.close()

#global variables
levelSelect = Scale(root)
levelMode = 0
raceNum = -1
subraceDrop = ttk.Combobox()
subraces = ["Select Race First"]
subraceNum = -1
ageRange = "Age Range"
ageLabel = Label()
ageMin = 0
ageMax = 0
ageRandReturn = " "
avgHeight = "__ feet "
heightLabel = Label()
heightMod = "__ inches"
heightModLabel = Label()
randAbilities = "__,  __,  __,  __,  __,  __"
randomAbilitiesLabel = Label()
abilityScore = [0]*6
abilityMod = [0]*6
abilityName = ["STR", "DEX", "CON", "INT", "WIS", "CHA"]
abilityDisplay = ["__           __           __","__           __           __","__           __           __","__           __           __","__           __           __","__           __           __"]
scoreDrop = [ttk.Combobox(), ttk.Combobox(), ttk.Combobox(), ttk.Combobox(), ttk.Combobox(), ttk.Combobox()]
abilityLabel = [Label(), Label(), Label(), Label(), Label(), Label()]
classNum = -1
mageLabel = Label()
slotNumbering = ['1st', '2nd', '3rd', '4th', '5th', '6th', '7th', '8th', '9th']
subclassDrop = ttk.Combobox()
subclassNum = -1
health = [0]*20
healthString = " "
healthLabel = Label(root)
backgroundNum = -1
exportButton = Button()

#Begin Function Definitions

def setLevelMode(mode):
    global levelMode
    levelMode = mode

def raceSet(event): #sets player race to dropdown box selection, displays average height and age range, applies racial ability modifiers
    global raceInfo
    global raceData
    global raceNum
    global subraceNum
    global subraces
    global ageRange
    global ageLabel
    global ageMax
    global ageMin
    global subraceDrop
    global avgHeight
    global heightLabel
    raceNum = 0
    subraceNum =-1
    while (raceDrop.get() != raceData[raceNum]['name']) and (raceNum <= len(raceData)):
        raceNum += 1
    if 'subrace' in raceData[raceNum]: subraces = list(map(lambda x : x['name'], raceData[raceNum]['subrace']))
    else: subraces = "No_Subraces"
    subraceDrop.destroy()
    subraceDrop = ttk.Combobox(root, value=subraces)
    subraceDrop.bind("<<ComboboxSelected>>", subraceSet)
    subraceDrop.grid(row=25, column=1, sticky=W)
    ageMax = raceData[raceNum]['ageMaximum']
    ageMin = raceData[raceNum]['ageMinimum']
    ageRange = str(int(ageMin)) + " to " + str(int(ageMax)) + " years"
    ageLabel = Label(root, text=ageRange).grid(row=32, column=2)
    avgHeight = str(raceData[raceNum]['heightAvg']) + "  feet"
    heightLabel.destroy()
    heightLabel = Label(root, text=avgHeight)
    heightLabel.grid(row=35, column=1)
    abilitySet(0)

def abilityString(abilities):
    out = "STR: " + str(abilities[0]) + "          INT: " + str(abilities[3]) + "\nDEX: " + str(abilities[1]) + "          WIS: " + str(abilities[4]) + "\nCON: " + str(abilities[2]) + "          CHA: " + str(abilities[5])
    return out

def raceInfoDisplay(): #displays selected race info when clicking button
    global raceData
    global raceNum
    global subraceNum
    if raceNum >= 0:
        raceInfoWindow = Toplevel()
        raceInfoWindow.title('DDCC Race Info')
        label = Label(raceInfoWindow, text=raceData[raceNum]['name'], justify=LEFT, font=(18)).grid(row=1, column=0, columnspan=2)
        label = Label(raceInfoWindow, text="Ability Bonus: ", anchor=NE, height=3).grid(row=2, column=0, sticky=E)
        label = Label(raceInfoWindow, text=abilityString(raceData[raceNum]['abilityBonus'])).grid(row=2, column=1, sticky=W)
        label = Label(raceInfoWindow, text="Age Minimum: ").grid(row=3, column=0, sticky=E)
        label = Label(raceInfoWindow, text=(str(raceData[raceNum]['ageMinimum']) + " years")).grid(row=3, column=1, sticky=W)
        label = Label(raceInfoWindow, text="Age Maximum: ").grid(row=4, column=0, sticky=E)
        label = Label(raceInfoWindow, text=(str(raceData[raceNum]['ageMaximum']) + " years")).grid(row=4, column=1, sticky=W)
        label = Label(raceInfoWindow, text=("Average Height: ")).grid(row=5, column=0, sticky=E)
        label = Label(raceInfoWindow, text=(str(raceData[raceNum]['heightAvg']) + " feet")).grid(row=5, column=1, sticky=W)
        label = Label(raceInfoWindow, text="Size Class: ").grid(row=6, column=0, sticky=E)
        label = Label(raceInfoWindow, text=raceData[raceNum]['sizeClass']).grid(row=6, column=1, sticky=W)
        label = Label(raceInfoWindow, text="Base Speed: ").grid(row=7, column=0, sticky=E)
        label = Label(raceInfoWindow, text=(str(raceData[raceNum]['speedBase']) + " feet per round")).grid(row=7, column=1, sticky=W)
        if 'proficiencySkillSave' in raceData[raceNum]:
            label = Label(raceInfoWindow, text="Skill/Save Proficiencies: ").grid(row=8, column=0, sticky=E)
            label = Label(raceInfoWindow, text=raceData[raceNum]['proficiencySkillSave'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=8, column=1, sticky=W)
        if 'proficiencyArmsArmor' in raceData[raceNum]:
            label = Label(raceInfoWindow, text="Equipment Proficiencies: ").grid(row=9, column=0, sticky=E)
            label = Label(raceInfoWindow, text=raceData[raceNum]['proficiencyArmsArmor'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=9, column=1, sticky=W)
        if 'proficiencyTools' in raceData[raceNum]:
            label = Label(raceInfoWindow, text="Tool Proficiencies: ").grid(row=10, column=0, sticky=E)
            label = Label(raceInfoWindow, text=raceData[raceNum]['proficiencyTools'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=10, column=1, sticky=W)
        if 'languages' in raceData[raceNum]:
            label = Label(raceInfoWindow, text="Languages: ").grid(row=11, column=0, sticky=E)
            label = Label(raceInfoWindow, text=raceData[raceNum]['languages'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=11, column=1, sticky=W)
        if 'extras' in raceData[raceNum]:
            label = Label(raceInfoWindow, text="Additional Features: ").grid(row=12, column=0, sticky=E)
            label = Label(raceInfoWindow, text=raceData[raceNum]['extras'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=12, column=1, sticky=W)

def subraceSet(event):
    global subraceNum
    global raceNum
    global raceData
    subraceNum = 0
    if 'subrace' in raceData[raceNum]: 
        while (subraceDrop.get() != raceData[raceNum]['subrace'][subraceNum]['name']) and (subraceNum <= len(raceData[raceNum]['subrace'])):
            subraceNum += 1
    abilitySet(0)

def subraceInfoDisplay():
    global raceData
    global raceNum
    global subraceNum
    if 'subrace' in raceData[raceNum] and subraceNum >= 0:
        subraceInfoWindow = Toplevel()
        subraceInfoWindow.title('DDCC Sub-Race Info')
        label = Label(subraceInfoWindow, text=(raceData[raceNum]['name'] + ",  " + raceData[raceNum]['subrace'][subraceNum]['name']), justify=LEFT, font=(18)).grid(row=1, column=0, columnspan=2)
        label = Label(subraceInfoWindow, text="Ability Bonus: ", anchor=NE, height=3).grid(row=2, column=0, sticky=E)
        label = Label(subraceInfoWindow, text=abilityString(raceData[raceNum]['subrace'][subraceNum]['abilityBonus'])).grid(row=2, column=1, sticky=W)
        if 'proficiencySkillSave' in raceData[raceNum]['subrace'][subraceNum]:
            label = Label(subraceInfoWindow, text="Skill/Save Proficiencies: ").grid(row=8, column=0, sticky=E)
            label = Label(subraceInfoWindow, text=raceData[raceNum]['subrace'][subraceNum]['proficiencySkillSave'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=8, column=1, sticky=W)
        if 'proficiencyArmsArmor' in raceData[raceNum]['subrace'][subraceNum]:
            label = Label(subraceInfoWindow, text="Equipment Proficiencies: ").grid(row=9, column=0, sticky=E)
            label = Label(subraceInfoWindow, text=raceData[raceNum]['subrace'][subraceNum]['proficiencyArmsArmor'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=9, column=1, sticky=W)
        if 'proficiencyTools' in raceData[raceNum]['subrace'][subraceNum]:
            label = Label(subraceInfoWindow, text="Tool Proficiencies: ").grid(row=10, column=0, sticky=E)
            label = Label(subraceInfoWindow, text=raceData[raceNum]['subrace'][subraceNum]['proficiencyTools'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=10, column=1, sticky=W)
        if 'languages' in raceData[raceNum]['subrace'][subraceNum]:
            label = Label(subraceInfoWindow, text="Languages: ").grid(row=11, column=0, sticky=E)
            label = Label(subraceInfoWindow, text=raceData[raceNum]['subrace'][subraceNum]['languages'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=11, column=1, sticky=W)
        if 'extras' in raceData[raceNum]['subrace'][subraceNum]:
            label = Label(subraceInfoWindow, text="Additional Features: ").grid(row=12, column=0, sticky=E)
            label = Label(subraceInfoWindow, text=raceData[raceNum]['subrace'][subraceNum]['extras'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=12, column=1, sticky=W)

def ageRand():
    global ageMin
    global ageMax
    age = random.randint(ageMin, ageMax)
    ageRandReturn = str(age) + " years  "
    ageRandReturnLabel = Label(root, text=ageRandReturn).grid(row=30, column=1)

def randHeight():
    global heightMod
    global heightModLabel
    modifier = random.randint(-12, 12)
    heightMod = str(modifier) + " inches"
    heightModLabel.destroy()
    heightModLabel = Label(root, text=heightMod)
    heightModLabel.grid(row=36, column=3, sticky=W)

def abilitiesRNG(RNGmode):
    global randAbilities
    global randomAbilitiesLabel
    score = [0, 0, 0, 0, 0, 0]
    if RNGmode == 1:
        for x in range (0, 6): #1-20 random ability
            score[x] = random.randint(1, 20)
    if RNGmode == 2:
        for x in range (0, 6): #10-20 random ability
            score[x] = random.randint(10, 20)
    if RNGmode == 3:
        for x in range (0, 6): #4d6-lowest
            value = [0, 0, 0, 0]
            for y in range (0, 3):
                value[y] = random.randint(1, 6)  
            value.sort()
            score[x] = value[1] + value[2] + value[3]
    randAbilities = abilityFormat(score[0])
    for x in range (1, 6):
        randAbilities = randAbilities + ",  " + abilityFormat(score[x])
    randomAbilitiesLabel = Label(root, text=randAbilities, justify=LEFT).grid(row=41, column=2, columnspan=3, sticky=W)

def fullRandAbilities():
    abilitiesRNG(1)

def halfRandAbilities():
    abilitiesRNG(2)

def standardAbilities():
    abilitiesRNG(3)

def abilitySet(event):
    global raceData
    global raceNum
    global subraceNum
    global abilityLabel
    global abilityScore
    global abilityMod
    for x in range (0, 6):
        score = 0
        if scoreDrop[x].get(): score = int(scoreDrop[x].get())
        raceBonus = 0
        if raceDrop.get(): raceBonus = raceData[raceNum]['abilityBonus'][x]
        subraceBonus = 0
        if 'subrace' in raceData[raceNum] and subraceDrop.get(): subraceBonus = raceData[raceNum]['subrace'][subraceNum]['abilityBonus'][x]
        abilityScore[x] = raceBonus + subraceBonus + score
        abilityMod[x] = int(math.floor((raceBonus + subraceBonus + score - 10) / 2))
        abilityDisplay[x] = abilityFormat(raceBonus + subraceBonus) + "          " + abilityFormat(abilityScore[x])  + "          " + abilityFormat(abilityMod[x])
        abilityLabel[x] = Label(root, text=abilityDisplay[x], justify=LEFT).grid(row=(46 + x), column=2, sticky=W)

def abilityFormat(value):
    out = str(value)
    if (value < 10) and (value >= 0): out = "  " + str(value)
    return(out)

def classSet(event):
    global classInfo
    global subclassInfo
    global subclassNum
    global classNum
    global classData
    global subclassDrop
    global mageLabel
    classNum = 0
    subclassNum = -1
    while (classDrop.get() != classData[classNum]['name']) and (classNum <= len(classData)):
        classNum += 1
    mageLabel.destroy()
    if 'spellSlots' in classData[classNum]:
        mageLabel = Label(root, text="Class is a mage")
        mageLabel.grid(row=61, column=1)
    if 'subClass' in classData[classNum]: subclasses = list(map(lambda x : x['name'], classData[classNum]['subClass']))
    else: subclasses = "No_Subclasses"
    subclassDrop.destroy()
    subclassDrop = ttk.Combobox(root, value=subclasses)
    subclassDrop.bind("<<ComboboxSelected>>", subclassSet)
    subclassDrop.grid(row=65, column=1, sticky=W)

def classInfoDisplay():
    global classData
    global classNum
    if classDrop.get():
        classInfoWindow = Toplevel(root)
        classInfoWindow.title('DDCC Class Info')
        label = Label(classInfoWindow, width=50).grid(row=0, column=1)
        label = Label(classInfoWindow, text=classData[classNum]['name'], justify=LEFT, font=(18)).grid(row=1, column=0, columnspan=2)
        label = Label(classInfoWindow, text="Hit Die: ").grid(row=3, column=0, sticky=E)
        label = Label(classInfoWindow, text=("d" + str(classData[classNum]['hitDie']))).grid(row=3, column=1, sticky=W)
        label = Label(classInfoWindow, text="Skill/Save Proficiencies: ").grid(row=4, column=0, sticky=E)
        label = Label(classInfoWindow, text=classData[classNum]['proficiencySkillSave'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=4, column=1, sticky=W)
        if 'proficiencyArmsArmor' in classData[classNum]:
            label = Label(classInfoWindow, text="Equipment Proficiencies: ").grid(row=5, column=0, sticky=E)
            label = Label(classInfoWindow, text=classData[classNum]['proficiencyArmsArmor'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=5, column=1, sticky=W)
        if 'proficiencyTools' in classData[classNum]:
            label = Label(classInfoWindow, text="Tool Proficiencies: ").grid(row=6, column=0, sticky=E)
            label = Label(classInfoWindow, text=classData[classNum]['proficiencyTools'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=6, column=1, sticky=W)
        if 'languages' in classData[classNum]:
            label = Label(classInfoWindow, text="Languages: ").grid(row=7, column=0, sticky=E)
            label = Label(classInfoWindow, text=classData[classNum]['languages'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=7, column=1, sticky=W)
        for x in range (1, 21):
            if ('lvl' + str(x)) in classData[classNum]:
                label = Label(classInfoWindow, text=("Lvl " + str(x) +":")).grid(row=11+x, column=0, sticky=E)
                label = Label(classInfoWindow, text=classData[classNum][('lvl' + str(x))], anchor=W , borderwidth=2, relief="ridge", width=75, wraplength=525, justify=LEFT)
                label.grid(row=11+x, column=1, sticky=W)
                if int(label.winfo_reqheight()) > 80:
                    label.destroy()
                    text = Text(classInfoWindow, height=4, bd=2, relief="ridge", width=87, wrap=WORD, font="TkDefaultFont")
                    text.grid(row=11+x, column=1, sticky=W)
                    text.insert(END, classData[classNum][('lvl' + str(x))])
                    text.config(state=DISABLED)

def spellSlotDisplay():
    global classData
    global classNum
    global slotNumbering
    if classDrop.get() and 'spellSlots' in classData[classNum]:
        spellInfoWindow = Toplevel(root)
        spellInfoWindow.title('DDCC Spell Slot Info')
        for x in range (0, 20):
            label = Label(spellInfoWindow, text='Level ' + str(x+1)).grid(row=10+x, column=0, sticky=E)
        label = Label(spellInfoWindow, text=(classData[classNum]['name'] + ' Spell Slots'), font=(18)).grid(row=1, column=0, columnspan=12)
        if 'cantripsKnown' in classData[classNum]:
            label =Label(spellInfoWindow, text="Cantrips Known", wraplength=50).grid(row=5, column=2)
            for x in range (0, 20):
                label = Label(spellInfoWindow, text=str(classData[classNum]['cantripsKnown'][x]), width=3, borderwidth=2, relief="ridge").grid(row=10+x, column=2)
        if 'spellsKnown' in classData[classNum]:
            label =Label(spellInfoWindow, text="Spells Known", wraplength=50).grid(row=5, column=3)
            for x in range (0, 20):
                label = Label(spellInfoWindow, text=str(classData[classNum]['spellsKnown'][x]), width=3, borderwidth=2, relief="ridge").grid(row=10+x, column=3)
        for y in range (0, 9):
            label = Label(spellInfoWindow, text=slotNumbering[y]).grid(row=5, column=10+y)
            for x in range (0, 20):
                if str(classData[classNum]['spellSlots'][y][x]) == '0': out = '--'
                else: out = str(classData[classNum]['spellSlots'][y][x])
                label = Label(spellInfoWindow, text=out, borderwidth=2, relief="ridge", width=3).grid(row=10+x, column=10+y)

def subclassSet(event):
    global subclassInfo
    global subclassNum
    global classNum
    global classData
    subclassNum = 0
    if 'subClass' in classData[classNum]:
        while (subclassDrop.get() != classData[classNum]['subClass'][subclassNum]['name']) and (subclassNum <= len(classData[classNum]['subClass'])):
            subclassNum += 1

def subclassInfoDisplay():
    global subclassNum
    if subclassDrop.get() and (not subclassDrop.get() == "No_Subclasses"):
        subclassInfoWindow = Toplevel(root)
        subclassInfoWindow.title('DDCC Sub-Class Info')
        label = Label(subclassInfoWindow, width=50).grid(row=0, column=1)
        label = Label(subclassInfoWindow, text=classData[classNum]['subClass'][subclassNum]['name'], justify=LEFT, font=(18)).grid(row=1, column=0, columnspan=2)
        for x in range (1, 21):
            if ('lvl' + str(x)) in classData[classNum]['subClass'][subclassNum]:
                label = Label(subclassInfoWindow, text="Lvl " + str(x) + ":").grid(row=11+x, column=0, sticky=E)
                label = Label(subclassInfoWindow, text=classData[classNum]['subClass'][subclassNum][('lvl' + str(x))], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=11+x, column=1, sticky=W)

def genHealth(mode):
    global health
    global healthString
    global healthLabel
    global levelMode
    global levelSelect
    if classDrop.get():
        die = classData[classNum]['hitDie']
        for x in range (0, 20):
            if mode == 1:
                if x > 0: health[x] = health[x-1] + random.randint(1, die) + abilityMod[2]
                elif x == 0: health[x] = die + abilityMod[2]
            elif mode == 2:
                if x > 0: health[x] = health[x-1] + int(die / 2) + 1 + abilityMod[2]
                elif x == 0: health[x] = die + abilityMod[2]
            elif mode == 3:
                if x > 0: health[x] = health[x-1] + die + abilityMod[2]
                elif x == 0: health[x] = die + abilityMod[2]
        if levelMode == 1:
            healthString = str(health[levelSelect.get()-1]) + " HP"
        else:
            temp = 1
            for x in range (0, 20):
                if x == 0: healthString = "Lvl " + str(x+1) + ": " + str(health[x]) + " HP     "
                else: healthString += "Lvl " + str(x+1) + ": " + str(health[x]) + " HP    "
                if temp == 4:
                    healthString += "\n"
                    temp =0
                temp += 1
        healthLabel.destroy()
        healthLabel = Label(root, text=healthString, justify=LEFT, height=5, anchor=N)
        healthLabel.grid(row=71, column=1, columnspan=3, sticky=W)

def rollHealth():
    genHealth(1)

def avgHealth():
    genHealth(2)

def maxHealth():
    genHealth(3)

def backgroundSet(event):
    global backgroundData
    global backgroundNum
    backgroundNum = 0
    while (backgroundDrop.get() != backgroundData[backgroundNum]['name']) and (backgroundNum <= len(backgroundData)):
        backgroundNum += 1
    
def backgroundInfoDisplay():
    global backgroundData
    global backgroundNum
    if backgroundDrop.get():
        backgroundInfoWindow = Toplevel(root)
        backgroundInfoWindow.title('DDCC Background Info')
        label = Label(backgroundInfoWindow, text=backgroundData[backgroundNum]['name'], justify=LEFT, font=(18)).grid(row=1, column=0, columnspan=2)
        label = Label(backgroundInfoWindow, text="Feature: ").grid(row=2, column=0, sticky=E)
        label = Label(backgroundInfoWindow, text=backgroundData[backgroundNum]['feature'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=2, column=1, sticky=W)
        if 'proficiencySkillSave' in backgroundData[backgroundNum]:
            label = Label(backgroundInfoWindow, text="Skill/Save Proficiencies: ").grid(row=8, column=0, sticky=E)
            label = Label(backgroundInfoWindow, text=backgroundData[backgroundNum]['proficiencySkillSave'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=8, column=1, sticky=W)
        if 'proficiencyArmsArmor' in backgroundData[backgroundNum]:
            label = Label(backgroundInfoWindow, text="Equipment Proficiencies: ").grid(row=9, column=0, sticky=E)
            label = Label(backgroundInfoWindow, text=backgroundData[backgroundNum]['proficiencyArmsArmor'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=9, column=1, sticky=W)
        if 'proficiencyTools' in backgroundData[backgroundNum]:
            label = Label(backgroundInfoWindow, text="Tool Proficiencies: ").grid(row=10, column=0, sticky=E)
            label = Label(backgroundInfoWindow, text=backgroundData[backgroundNum]['proficiencyTools'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=10, column=1, sticky=W)
        if 'languages' in backgroundData[backgroundNum]:
            label = Label(backgroundInfoWindow, text="Languages: ").grid(row=11, column=0, sticky=E)
            label = Label(backgroundInfoWindow, text=backgroundData[backgroundNum]['languages'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=11, column=1, sticky=W)
        if 'extras' in backgroundData[backgroundNum]:
            label = Label(backgroundInfoWindow, text="Additional Features: ").grid(row=12, column=0, sticky=E)
            label = Label(backgroundInfoWindow, text=backgroundData[backgroundNum]['extras'], anchor=W , borderwidth=2, relief="ridge", width=50, wraplength=350, justify=LEFT).grid(row=12, column=1, sticky=W)

def checkForm(): #check completion of all fields to unlock export button
    global exportButton
    global nameEntry
    if not nameEntry.get(): messagebox.showinfo("Missing Requirements", "Please Enter Character Name")
    elif not raceDrop.get(): messagebox.showinfo("Missing Requirements", "Please Enter Character Race")
    elif (not subraceDrop.get()) and ('subrace' in raceData[raceNum]): messagebox.showinfo("Missing Requirements", "Please Enter Character Subrace")
    elif not scoreDrop[0].get(): messagebox.showinfo("Missing Requirements", "Please Enter Character Strength")
    elif not scoreDrop[1].get(): messagebox.showinfo("Missing Requirements", "Please Enter Character Dexterity")
    elif not scoreDrop[2].get(): messagebox.showinfo("Missing Requirements", "Please Enter Character Constitution")
    elif not scoreDrop[3].get(): messagebox.showinfo("Missing Requirements", "Please Enter Character Intelligence")
    elif not scoreDrop[4].get(): messagebox.showinfo("Missing Requirements", "Please Enter Character Wisdom")
    elif not scoreDrop[5].get(): messagebox.showinfo("Missing Requirements", "Please Enter Character Charisma")
    elif not classDrop.get(): messagebox.showinfo("Missing Requirements", "Please Enter Character Class")
    elif (not subclassDrop.get()) and ('subClass' in classData[classNum]): messagebox.showinfo("Missing Requirements", "Please Enter Character Subclass")
    elif health[0] <= 0: messagebox.showinfo("Missing Requirements", "Please Generate Character Health")
    elif not backgroundDrop.get(): messagebox.showinfo("Missing Requirements", "Please Enter Character Background")
    else:
        exportButton["state"] = "normal"

def export():
    global exportButton
    global slotNumbering
    temp = " "
    doc = docx.Document()
    doc.add_heading(nameEntry.get(), 0)
    #basic information
    if levelMode == 1: temp += "Level " + str(levelSelect.get()) + " "
    if 'subrace' in raceData[raceNum]: temp += raceData[raceNum]['subrace'][subraceNum]['name'] + " "
    temp += raceData[raceNum]['name'] + " " + classData[classNum]['name']
    if 'subClass' in classData[classNum]: temp += " (" + classData[classNum]['subClass'][subclassNum]['name'] + ")"
    temp += " - " + backgroundData[backgroundNum]['name'] + " Background"
    doc.add_paragraph(temp)
    #Age and Height, size class
    height = ((raceData[raceNum]['heightAvg'] * 12) + heightScale.get())
    feet = int(height // 12)
    inches = int(height % 12)
    doc.add_paragraph("Age: " + str(ageEntry.get()) + " years          Height: " + str(feet) + " Feet, " + str(inches) + " inches          Size: " + raceData[raceNum]['sizeClass'])
    #Initiative, speed, passive wisdom
    doc.add_paragraph("Initiative Bonus: " + str(abilityMod[1]) + "          Speed: " + str(raceData[raceNum]['speedBase']) + "          Passive Perception: " + str(abilityMod[4] + 8))
    #health
    if levelMode == 1: doc.add_paragraph("Max Health: " + str(health[levelSelect.get() - 1]) + " HP")
    else:
        doc.add_heading("Max Health by Level", 1)
        col = 0
        healthTable = doc.add_table(rows = 1, cols = 5)
        rowCells = healthTable.rows[0].cells
        for x in range (0, 20):
            rowCells[col].text = "Lvl " + str(x+1) + ": " + str(health[x])
            col += 1
            if col == 5:
                col = 0
                rowCells = healthTable.add_row().cells
    #Ability scores and modifiers
    doc.add_heading("Ability Scores", 1)
    abilityTable = doc.add_table(rows = 1, cols = 3)
    rowCells = abilityTable.rows[0].cells
    for x in range (0, 6):
        rowCells[0].text = abilityName[x]
        rowCells[1].text = str(abilityScore[x])
        if abilityMod[x] > 0: rowCells[2].text = "+" + str(abilityMod[x])
        else: rowCells[2].text = str(abilityMod[x])
        if not (x >= 5): rowCells = abilityTable.add_row().cells
    #Proficiencies
    doc.add_heading("Proficiencies", 1)
    doc.add_heading("Skills and Saves", 2)
    doc.add_paragraph(classData[classNum]['proficiencySkillSave'])
    if 'proficiencySkillSave' in raceData[raceNum]: doc.add_paragraph(raceData[raceNum]['proficiencySkillSave'])
    if 'subrace' in raceData[raceNum]: 
        if 'proficiencySkillSave' in raceData[raceNum]['subrace'][subraceNum]: doc.add_paragraph(raceData[raceNum]['subrace'][subraceNum]['proficiencySkillSave'])
    if 'proficiencySkillSave' in backgroundData[backgroundNum]: doc.add_paragraph(backgroundData[backgroundNum]['proficiencySkillSave'])
    if ('proficiencyArmsArmor' in raceData[raceNum]) or ('proficiencyArmsArmor' in backgroundData[backgroundNum]) or ('proficiencyArmsArmor' in classData[classNum]):
        doc.add_heading("Arms and Armor", 2)
        if 'proficiencyArmsArmor' in classData[classNum]: doc.add_paragraph(classData[classNum]['proficiencyArmsArmor'])
        if 'proficiencyArmsArmor' in raceData[raceNum]: doc.add_paragraph(raceData[raceNum]['proficiencyArmsArmor'])
        if 'subrace' in raceData[raceNum]: 
            if 'proficiencyArmsArmor' in raceData[raceNum]['subrace'][subraceNum]: doc.add_paragraph(raceData[raceNum]['subrace'][subraceNum]['proficiencyArmsArmor'])
        if 'proficiencyArmsArmor' in backgroundData[backgroundNum]: doc.add_paragraph(backgroundData[backgroundNum]['proficiencyArmsArmor'])
    if ('proficiencyTools' in raceData[raceNum]) or ('proficiencyTools' in backgroundData[backgroundNum]) or ('proficiencyTools' in classData[classNum]):
        doc.add_heading("Tools", 2)
        if 'proficiencyTools' in classData[classNum]: doc.add_paragraph(classData[classNum]['proficiencyTools'])
        if 'proficiencyTools' in raceData[raceNum]: doc.add_paragraph(raceData[raceNum]['proficiencyTools'])
        if 'subrace' in raceData[raceNum]: 
            if 'proficiencyTools' in raceData[raceNum]['subrace'][subraceNum]: doc.add_paragraph(raceData[raceNum]['subrace'][subraceNum]['proficiencyTools'])
        if 'proficiencyTools' in backgroundData[backgroundNum]: doc.add_paragraph(backgroundData[backgroundNum]['proficiencyTools'])
    #Spell Slots
    if 'spellSlots' in classData[classNum]:
        doc.add_heading("Spell Slots", 1)
        spellTable = doc.add_table(rows = 1, cols = 12)
        spellTable.style= 'Medium List 2'
        rowCells = spellTable.rows[0].cells
        if 'cantripsKnown' in classData[classNum]:
            rowCells[1].text = 'Cantrips Known'
        if 'spellsKnown' in classData[classNum]:
            rowCells[2].text = 'Spells Known'
        for x in range (0, 9):
            rowCells[3+x].text = slotNumbering[x]
        if not (levelMode == 1): #chart of all spell slots for all levels
            for x in range (0, 20):
                rowCells = spellTable.add_row().cells
                rowCells[0].text = "level " + str(x+1)
                if 'cantripsKnown' in classData[classNum]:
                    rowCells[1].text = str(classData[classNum]['cantripsKnown'][x])
                if 'spellsKnown' in classData[classNum]:
                    rowCells[2].text = str(classData[classNum]['spellsKnown'][x])
                for y in range (0, 9):
                    if str(classData[classNum]['spellSlots'][y][x]) == '0': out = '--'
                    else: out = str(classData[classNum]['spellSlots'][y][x])
                    rowCells[3+y].text = out
        else: #chart for spell slots for specific level
            rowCells = spellTable.add_row().cells
            rowCells[0].text = "level " + str(levelSelect.get())
            if 'cantripsKnown' in classData[classNum]:
                rowCells[1].text = str(classData[classNum]['cantripsKnown'][levelSelect.get()])
            if 'spellsKnown' in classData[classNum]:
                 rowCells[2].text = str(classData[classNum]['spellsKnown'][levelSelect.get()])
            for y in range (0, 9):
                if str(classData[classNum]['spellSlots'][y][levelSelect.get()]) == '0': out = '--'
                else: out = str(classData[classNum]['spellSlots'][y][levelSelect.get()])
                rowCells[3+y].text = out
    #features
    doc.add_heading("Features and Abilities", 2)
    if 'extras' in classData[classNum]: doc.add_paragraph(classData[classNum]['extras'])
    if 'extras' in raceData[raceNum]: doc.add_paragraph(raceData[raceNum]['extras'])
    if 'subrace' in raceData[raceNum]: 
        if 'extras' in raceData[raceNum]['subrace'][subraceNum]: doc.add_paragraph(raceData[raceNum]['subrace'][subraceNum]['extras'])
    if 'extras' in backgroundData[backgroundNum]: doc.add_paragraph(backgroundData[backgroundNum]['extras'])
    doc.add_paragraph(backgroundData[backgroundNum]['feature'])
    doc.add_paragraph(classData[classNum]['lvl1'])
    if 'lvl1' in classData[classNum]['subClass'][subclassNum]: doc.add_paragraph("--Lvl 1: " + classData[classNum]['subClass'][subclassNum]['lvl1'])
    for x in range (1, 21):
        if (not (levelMode == 1)) or levelSelect.get() >= x:
            if 'lvl' + str(x) in classData[classNum]: doc.add_paragraph("Lvl " + str(x) + ": " + classData[classNum]['lvl' + str(x)])
            if 'lvl' + str(x) in classData[classNum]['subClass'][subclassNum]: doc.add_paragraph("--Lvl " + str(x) + ": " + classData[classNum]['subClass'][subclassNum]['lvl' + str(x)])
    #languages
    doc.add_heading("Languages", 2)
    if 'languages' in classData[classNum]: doc.add_paragraph(classData[classNum]['languages'])
    if 'languages' in raceData[raceNum]: doc.add_paragraph(raceData[raceNum]['languages'])
    if 'subrace' in raceData[raceNum]: 
        if 'languages' in raceData[raceNum]['subrace'][subraceNum]: doc.add_paragraph(raceData[raceNum]['subrace'][subraceNum]['languages'])
    if 'languages' in backgroundData[backgroundNum]: doc.add_paragraph(backgroundData[backgroundNum]['languages'])
    #Final
    doc.save(nameEntry.get() + ".docx")
    exportButton["state"] = DISABLED

#Formatting labels for rows 0 and 99, and column 99
formatLabel = Label(root, width=20).grid(row=0, column=1)
formatLabel = Label(root, text=" ", width=15).grid(row=0, column=2)
formatLabel = Label(root, text=" ", width=10).grid(row=0, column=3)


#Name Entry Block
label = Label(root, text="Character Name: ").grid(row=1, column=0, sticky=E)
nameEntry = Entry(root)
nameEntry.grid(row=1, column=1, columnspan=2, sticky=W)

#Level select and mode block
levelModeSet = IntVar()
label = Label(root, text="Select Level Mode: ").grid(row=5, column=0, sticky=E)
Radiobutton(root, text="Set", variable=levelModeSet, value=1, command=lambda: setLevelMode(levelModeSet.get())).grid(row=5, column=1)
Radiobutton(root, text="Map", variable=levelModeSet, value=2, command=lambda: setLevelMode(levelModeSet.get())).grid(row=5, column=2)
levelSelectPrompt = Label(root, text="Enter Level: ").grid(row=6, column=1, sticky=E)
levelSelect = Scale(root, from_=1, to=20, orient=HORIZONTAL)
levelSelect.grid(row=6, column=2)

#Race Select Block 
label = Label(root, text="Race: ").grid(row=20, column=0, sticky=E)
races = list(map(lambda x : x['name'], raceData))
raceDrop = ttk.Combobox(root, value=races)
raceDrop.bind("<<ComboboxSelected>>", raceSet)
raceDrop.grid(row=20, column=1, sticky=W)
raceDisplayButton = Button(root, text="Info", command=raceInfoDisplay).grid(row=20, column=2)

#Subrace Select Block
label = Label(root, text="Subrace: ").grid(row=25, column=0, sticky=E)
if 'subrace' in raceData[raceNum]: subraces = list(map(lambda x : x['name'], raceData[raceNum]['subrace']))
else: subraces = "No_Subraces"
subraceDrop = ttk.Combobox(root, value=subraces)
subraceDrop.bind("<<ComboboxSelected>>", subraceSet)
subraceDrop.grid(row=25, column=1, sticky=W)
subraceDisplayButton = Button(root, text="Info", command=subraceInfoDisplay).grid(row=25, column=2)

#Age Select and Mode block
label = Label(root, text="Generate Age: ").grid(row=30, column=0, sticky=E)
ageRandButton = Button(root, text="Random", command=ageRand).grid(row=30, column=2)
ageRandReturnLabel = Label(root, text=ageRandReturn).grid(row=30, column=1)
label = Label(root, text="Age: ").grid(row=32, column=0, sticky=E)
ageEntry = Entry(root)
ageEntry.grid(row=32, column=1, sticky=W)
ageLabel = Label(root, text=ageRange).grid(row=32, column=2)

#Height select and mode block
label = Label(root, text="Height: ").grid(row=35, column=0, sticky=E)
heightLabel = Label(root, text=avgHeight)
heightLabel.grid(row=35, column=1)
heightScale = Scale(root, from_=-12, to=12, orient=HORIZONTAL)
heightScale.grid(row=35, column=2)
label = Label(root, text="inches").grid(row=35, column=3, sticky=W)
label = Label(root, text="Random height modifier: ").grid(row=36, column=1, sticky=E)
button = Button(root, text="Generate", command=randHeight).grid(row=36, column=2)
heightModLabel = Label(root, text=heightMod)
heightModLabel.grid(row=36, column=3, sticky=W)

#Ability Bonus Selection and Mode block
label = Label(root, text="Ability Score RNG: ").grid(row=40, column=0)
button = Button(root, text="1-20", command=fullRandAbilities).grid(row=40, column=1)
button = Button(root, text="10-20", command=halfRandAbilities).grid(row=40, column=2)
button = Button(root, text="4d6-low", command=standardAbilities).grid(row=40, column=3)
label = Label(root, text="Generated Scores: ").grid(row=41, column=1, sticky=E)
randomAbilitiesLabel = Label(root, text=randAbilities, justify=LEFT).grid(row=41, column=2, columnspan=3, sticky=W)
label = Label(root, text="Selection").grid(row=45, column=1)
label = Label(root, text="Race     Score     Bonus").grid(row=45, column=2, sticky=W)
abilityRange = [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20]
for x in range (0, 6):
    label = Label(root, text=abilityName[x]).grid(row=(46+x), column=0, sticky=E)
    scoreDrop[x] = ttk.Combobox(root, value=abilityRange, width=2)
    scoreDrop[x].bind("<<ComboboxSelected>>", abilitySet)
    scoreDrop[x].grid(row=(46+x), column=1)
    abilityLabel[x] = Label(root, text=abilityDisplay[x]).grid(row=(46+x), column=2, sticky=W)

#Class select block
label = Label(root, text="Class: ").grid(row=60, column=0, sticky=E)
classes = list(map(lambda x : x['name'], classData))
classDrop = ttk.Combobox(root, value=classes)
classDrop.bind("<<ComboboxSelected>>", classSet)
classDrop.grid(row=60, column=1, sticky=W)
classDisplayButton = Button(root, text="Info", command=classInfoDisplay).grid(row=60, column=2)

#Spellcaster notification and display block
mageLabel = Label(root, text=" ")
mageLabel.grid(row=61, column=1)
button = Button(root, text="Slots", command=spellSlotDisplay).grid(row=61, column=2)

#Subclass select block
label = Label(root, text="Subclass: ").grid(row=65, column=0, sticky=E)
if 'subClass' in classData[classNum]: subclasses = list(map(lambda x : x['name'], classData[classNum]['subClass']))
else: subclasses = "No_Subclasses"
subclassDrop = ttk.Combobox(root, value=subclasses)
subclassDrop.bind("<<ComboboxSelected>>", subclassSet)
subclassDrop.grid(row=65, column=1, sticky=W)
subclassDisplayButton = Button(root, text="Info", command=subclassInfoDisplay).grid(row=65, column=2)

#Health and Initiative block
label = Label(root, text="Generate Health: ").grid(row=70, column=0, sticky=E)
button = Button(root, text="Roll", command=rollHealth).grid(row=70, column=1)
button = Button(root, text="Average", command=avgHealth).grid(row=70, column=2)
button = Button(root, text="Max", command=maxHealth).grid(row=70, column=3)
label = Label(root, text="Health: ", height=5, anchor=N).grid(row=71, column=0, sticky=E)
healthLabel = Label(root, text=healthString, anchor=N)
healthLabel.grid(row=71, column=1, columnspan=3, sticky=W)

#Background select block
label = Label(root, text="Background: ").grid(row=80, column=0, sticky=E)
backgrounds = list(map(lambda x : x['name'], backgroundData))
backgroundDrop = ttk.Combobox(root, value=backgrounds)
backgroundDrop.bind("<<ComboboxSelected>>", backgroundSet)
backgroundDrop.grid(row=80, column=1, sticky=W)
backgroundDisplayButton = Button(root, text="Info", command=backgroundInfoDisplay).grid(row=80, column=2)

#Completion buttons
checkButton = Button(root, text="Check", command=checkForm)
checkButton.grid(row=90, column=0)
exportButton = Button(root, text="Export", state=DISABLED, command=export)
exportButton.grid(row=90, column=1)
quitButton = Button(root, text="Quit", command=root.quit).grid(row=90, column=2)

root.mainloop()